"""
Generate VariantTools Architecture & Pipeline documentation as a Word .docx file.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    """Add a thin horizontal line paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading(doc, text, level, color_hex=None):
    h = doc.add_heading(text, level=level)
    if color_hex:
        for run in h.runs:
            run.font.color.rgb = RGBColor.from_string(color_hex)
    return h


def body(doc, text, bold_parts: dict = None):
    """Add a normal paragraph. bold_parts = {substring: True} to bold specific words."""
    p = doc.add_paragraph()
    if bold_parts:
        # Simple approach: split and rebuild
        remaining = text
        for phrase, _ in bold_parts.items():
            if phrase in remaining:
                before, _, after = remaining.partition(phrase)
                if before:
                    p.add_run(before)
                r = p.add_run(phrase)
                r.bold = True
                remaining = after
        if remaining:
            p.add_run(remaining)
    else:
        p.add_run(text)
    p.style = doc.styles["Normal"]
    return p


def code_block(doc, text):
    """Add a monospace code-style paragraph with grey background."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    # light-grey shading via pPr
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def simple_table(doc, headers, rows, header_color="1F497D"):
    """Create a styled table."""
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "FFFFFF" if ri % 2 == 0 else "F4F8FC"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    return table


# ── Document ────────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default paragraph font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ── Cover / Title ────────────────────────────────────────────────────────────

doc.add_paragraph()  # top spacer
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("VariantTools")
tr.bold = True
tr.font.size = Pt(28)
tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("Architecture, Pipeline & User Guide")
sr.font.size = Pt(16)
sr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"University of Minnesota — Lorenz Lab  |  v0.3.0  |  {datetime.date.today().strftime('%B %Y')}")
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# ── 1. Overview ─────────────────────────────────────────────────────────────

heading(doc, "1. Overview", 1, "1F497D")
body(doc,
     "VariantTools is a browser-based genomics variant processing suite designed for the "
     "soybean breeding program at UMN. It replaces a collection of ad-hoc R scripts and "
     "shell command workflows with a unified web application that lab members can use "
     "without needing to know the underlying command-line tools.")

doc.add_paragraph()
body(doc,
     "The application wraps industry-standard bioinformatics tools — bcftools, CrossMap, "
     "and custom Python parsing logic — behind a clean REST API and interactive web "
     "interface. All heavy computation runs asynchronously via a task queue, keeping the "
     "UI responsive even for multi-gigabyte VCF files.")

doc.add_paragraph()
body(doc, "Five processing modules are available:")

modules_overview = [
    ("VCF Stats & Assembly Detection", "Analyze VCF files, detect reference assembly from header, display bcftools stats"),
    ("Coordinate Liftover",            "Remap SNP positions between genome assemblies (A1→A6, A4→A6) using CrossMap"),
    ("Fix Reference (fixref)",         "Verify REF allele matches and flip/swap mismatches using bcftools +fixref"),
    ("Merge VCFs",                     "Combine multiple VCF files with optional multiallelic normalization"),
    ("VCF Generator",                  "Convert genotype dosage matrices (0/1/2) to standard VCF GT format"),
]
simple_table(doc, ["Module", "Purpose"], modules_overview)
doc.add_paragraph()

# ── 2. System Architecture ───────────────────────────────────────────────────

heading(doc, "2. System Architecture", 1, "1F497D")
body(doc,
     "VariantTools uses a three-tier architecture: a Next.js frontend communicates with a "
     "FastAPI backend over REST/HTTP, and the backend delegates long-running jobs to a "
     "Celery worker pool backed by Redis.")

doc.add_paragraph()

code_block(doc,
"""┌──────────────────────────────────────────────┐
│  Browser (Next.js 14 / React / TypeScript)   │
│  Tailwind CSS · Recharts · Lucide icons      │
│  http://localhost:3000                        │
└─────────────────────┬────────────────────────┘
                      │  HTTP / REST
┌─────────────────────▼────────────────────────┐
│  FastAPI Backend  (Python 3.11)              │
│  uvicorn · SQLModel · python-multipart       │
│  http://localhost:8000                        │
└──────┬──────────────────────┬────────────────┘
       │  Celery tasks        │  subprocess
┌──────▼──────┐    ┌──────────▼────────────────┐
│  Redis 7    │    │  bcftools · CrossMap       │
│  (broker +  │    │  bgzip · tabix            │
│   results)  │    └───────────────────────────┘
└─────────────┘
       │
┌──────▼──────────────────────────────────────┐
│  SQLite  (data/varianttools.db)             │
│  Job table · ChainFile table               │
└─────────────────────────────────────────────┘""")

doc.add_paragraph()

heading(doc, "2.1  Technology Stack", 2)
stack_rows = [
    ("Frontend",         "Next.js 14, React 18, TypeScript",   "App Router, server-side rendering"),
    ("UI / Styling",     "Tailwind CSS, Lucide React",          "Dark theme, responsive layout"),
    ("Charts",           "Recharts",                            "Substitution type & per-sample charts"),
    ("Backend API",      "FastAPI 0.111, Python 3.11",          "Async routes, automatic OpenAPI docs"),
    ("Task Queue",       "Celery 5.4 + Redis 7",                "Non-blocking long-running jobs"),
    ("Database",         "SQLite via SQLModel",                  "Job tracking, chain file registry"),
    ("File Storage",     "Local filesystem",                    "data/uploads/, data/outputs/, data/chains/"),
    ("Bioinformatics",   "bcftools, CrossMap 0.7.3, bgzip",     "Core genomics processing"),
    ("Containerisation", "Docker Compose (dev)",                "4-service stack; one command startup"),
    ("HPC Deployment",   "Singularity",                         "containers/varianttools.def for MSI"),
]
simple_table(doc, ["Layer", "Technology", "Notes"], stack_rows)
doc.add_paragraph()

# ── 3. Directory Structure ───────────────────────────────────────────────────

heading(doc, "3. Repository Layout", 1, "1F497D")
code_block(doc,
"""VariantTools/
├── backend/
│   ├── api/
│   │   ├── stats.py          # Module: VCF Stats & Assembly
│   │   ├── liftover.py       # Module: Coordinate Liftover
│   │   ├── fixref.py         # Module: Fix Reference
│   │   ├── merge.py          # Module: Merge VCFs
│   │   ├── generator.py      # Module: VCF Generator
│   │   └── slurm.py          # SLURM script generation
│   ├── core/
│   │   ├── bcftools.py       # async bcftools subprocess wrappers
│   │   ├── crossmap.py       # async CrossMap wrappers
│   │   ├── fixref.py         # async bcftools +fixref wrappers
│   │   ├── vcf_generator.py  # sync dosage-to-VCF conversion
│   │   ├── slurm.py          # SLURM script templates
│   │   └── file_manager.py   # upload/output path management
│   ├── db/
│   │   └── models.py         # Job + ChainFile SQLModel tables
│   ├── tasks/
│   │   ├── celery_app.py     # Celery configuration
│   │   └── vcf_tasks.py      # all Celery task definitions
│   └── main.py               # FastAPI app entry point
├── frontend/
│   ├── app/
│   │   ├── page.tsx          # Home dashboard
│   │   ├── stats/page.tsx
│   │   ├── liftover/page.tsx
│   │   ├── fixref/page.tsx
│   │   ├── merge/page.tsx
│   │   └── generator/page.tsx
│   └── components/
│       ├── NavBar.tsx
│       ├── FileUpload.tsx
│       └── StatsChart.tsx
├── containers/
│   ├── Dockerfile.backend
│   ├── Dockerfile.frontend
│   └── varianttools.def      # Singularity definition (MSI)
├── data/                     # runtime data (git-ignored)
│   ├── uploads/
│   ├── outputs/
│   └── chains/               # place .chain files here
├── docker-compose.yml
└── scripts/                  # original R / shell workflows""")
doc.add_paragraph()

# ── 4. Data Flow ─────────────────────────────────────────────────────────────

heading(doc, "4. Data Flow & Job Lifecycle", 1, "1F497D")
body(doc,
     "All processing modules follow the same asynchronous job pattern. This ensures the "
     "UI never blocks, even for operations that take minutes on large VCF files.")

doc.add_paragraph()
heading(doc, "4.1  Standard Job Flow", 2)

steps = [
    ("1", "Upload",    "User uploads file(s) via drag-and-drop. Frontend POSTs multipart/form-data to FastAPI."),
    ("2", "Save",      "Backend saves file to data/uploads/ with a UUID prefix. Returns file_id (the saved path)."),
    ("3", "Job create","A Job record is inserted into SQLite with status='pending'."),
    ("4", "Enqueue",   "A Celery task is dispatched to the Redis broker. Job record is updated with celery_task_id."),
    ("5", "Process",   "Celery worker picks up the task, runs the subprocess (bcftools / CrossMap), updates status to 'running'."),
    ("6", "Poll",      "Frontend polls GET /api/<module>/status/{job_id} every 2–3 seconds."),
    ("7", "Complete",  "Worker writes output to data/outputs/, stores result_json in DB, sets status='completed'."),
    ("8", "Download",  "User clicks Download; frontend calls GET /api/<module>/download/{job_id} which streams the file."),
]
simple_table(doc, ["Step", "Stage", "Description"], steps, header_color="2E75B6")
doc.add_paragraph()

heading(doc, "4.2  Synchronous Exception — Fix Reference Check", 2)
body(doc,
     "The 'Run Check' action in the Fix Reference module is the only synchronous operation. "
     "Because bcftools +fixref in check-only mode is fast (read-only, no output file), "
     "the FastAPI route runs it directly with asyncio.wait_for() and a 60-second timeout, "
     "returning the parsed stats immediately without creating a Celery job. "
     "The subsequent 'Run Fix' action follows the standard async job pattern.")
doc.add_paragraph()

# ── 5. Module Reference ───────────────────────────────────────────────────────

heading(doc, "5. Module Reference", 1, "1F497D")

# 5.1 Stats
heading(doc, "5.1  VCF Stats & Assembly Detection", 2)
body(doc,
     "This module provides two functions in a single upload step: instant assembly "
     "detection from the VCF header and full bcftools statistics via an async job.")
doc.add_paragraph()

heading(doc, "Assembly Detection", 3)
body(doc,
     "The VCF header is parsed synchronously on the server the moment the file is saved "
     "(before the stats job even starts). The parser scans ##reference= and ##contig= "
     "lines for known assembly name patterns:")
for pattern in [
    "Wm82.a6 / Wm82.a4 / Wm82.a2 / Wm82.a1 (soybean Williams 82)",
    "GRCh38 / hg38  and  GRCh37 / hg19 (human)",
    "B73 (maize), T2T-CHM13, GRCm38/mm10, GRCm39/mm39",
    "Unknown — displayed with a yellow badge if no pattern matches",
]:
    bullet(doc, pattern)
doc.add_paragraph()

heading(doc, "bcftools Stats", 3)
body(doc, "The async job runs:  bcftools stats <input.vcf>")
body(doc, "The stdout is parsed into three sections:")
for s in [
    "SN (Summary Numbers) — total SNPs, indels, MNPs, sample count, Ts/Tv ratio",
    "ST (Substitution Types) — counts and percentages for all 12 SNP types (A>C, A>G, ...)",
    "PSC (Per-Sample Counts) — hom-ref, het, hom-alt, missing per sample",
]:
    bullet(doc, s)
body(doc, "Results are displayed as stat cards + Recharts bar charts. Raw stats text is available for download.")
doc.add_paragraph()

api_rows_stats = [
    ("POST", "/api/stats/upload-and-analyze", "Upload VCF; returns assembly_info + job_id immediately"),
    ("GET",  "/api/stats/result/{job_id}",    "Poll for status and parsed stats"),
    ("GET",  "/api/stats/download/{job_id}",  "Download raw bcftools stats text"),
]
simple_table(doc, ["Method", "Endpoint", "Description"], api_rows_stats, header_color="1F6B3E")
doc.add_paragraph()

# 5.2 Liftover
heading(doc, "5.2  Coordinate Liftover", 2)
body(doc,
     "Remaps genomic coordinates (VCF or BED format) between genome assembly versions "
     "using CrossMap with FWD/REV chain files. This mirrors the CrossMap workflow "
     "documented in Liftover_a4_to_a6.txt and Liftover_A1_toA6.txt.")

doc.add_paragraph()
heading(doc, "Chain File Database", 3)
body(doc,
     "Chain files are managed through a database table (ChainFile). Four placeholder "
     "records are pre-seeded at startup for the Wm82 soybean assemblies:")
for cf in [
    "Wm82.a4 → Wm82.a6  FWD  (data/chains/Wm82_a4_to_a6_FWD.chain)",
    "Wm82.a4 → Wm82.a6  REV  (data/chains/Wm82_a4_to_a6_REV.chain)",
    "Wm82.a1 → Wm82.a6  FWD  (data/chains/Wm82_a1_to_a6_FWD.chain)",
    "Wm82.a1 → Wm82.a6  REV  (data/chains/Wm82_a1_to_a6_REV.chain)",
]:
    bullet(doc, cf)
body(doc,
     "Place the actual .chain files (generated from MUMmer + crossmap_delta_to_chain.pl) "
     "in the data/chains/ directory. The UI shows a 'file exists' badge per entry. "
     "Additional chain files can be uploaded directly from the Liftover page.")

doc.add_paragraph()
heading(doc, "CrossMap Execution", 3)
body(doc, "For VCF input:   CrossMap vcf <chain> <input.vcf> [ref.fa] <output.vcf>")
body(doc, "For BED input:   CrossMap bed <chain> <input.bed> <output.bed>")
body(doc,
     "Unmapped entries are written to <output>.unmap automatically by CrossMap. "
     "Both the mapped and unmapped files are available for download. "
     "The result card shows the mapped percentage with a colour-coded badge "
     "(green > 95%, yellow 80–95%, red < 80%).")
doc.add_paragraph()

api_rows_lo = [
    ("GET",  "/api/liftover/chain-files",         "List all chain files with file_exists status"),
    ("POST", "/api/liftover/chain-files/upload",  "Upload a custom chain file and register in DB"),
    ("POST", "/api/liftover/upload",              "Upload VCF or BED input file"),
    ("POST", "/api/liftover/submit",              "Submit liftover job; returns job_id"),
    ("GET",  "/api/liftover/status/{job_id}",     "Poll status + mapped/unmapped counts"),
    ("GET",  "/api/liftover/download/{job_id}",   "Download mapped output"),
    ("GET",  "/api/liftover/download-unmapped/{job_id}", "Download unmapped entries"),
]
simple_table(doc, ["Method", "Endpoint", "Description"], api_rows_lo, header_color="1F6B3E")
doc.add_paragraph()

# 5.3 Fix Reference
heading(doc, "5.3  Fix Reference (bcftools +fixref)", 2)
body(doc,
     "Checks and corrects REF allele mismatches between a VCF and a reference FASTA. "
     "This is essential after liftover, where the lifted coordinates may not match the "
     "target assembly's reference alleles.")
doc.add_paragraph()

heading(doc, "Two-step workflow", 3)

fix_steps = [
    ("Check (synchronous)",
     "bcftools +fixref <vcf> -- -f <ref.fa>",
     "Read-only. Returns ref_match %, mismatches, no output file. Result shown immediately."),
    ("Fix (async job)",
     "bcftools +fixref <vcf> -o <out.vcf> -- -f <ref.fa> -m flip",
     "Flips strand and swaps alleles to match reference. Returns fixed VCF for download."),
]
simple_table(doc, ["Step", "Command", "Description"], fix_steps, header_color="8B4513")
doc.add_paragraph()

body(doc,
     "The result card shows a before/after comparison table: total sites, ref match %, "
     "flipped count, swapped count, flip+swap count, and unresolved count. "
     "A 100% ref match after fix (as seen in the R script output) confirms success.")
body(doc,
     "The reference FASTA can be specified as a server-side absolute path "
     "(e.g. /scratch.global/vramasub/Gmax_880_v6.0.fa) or uploaded directly. "
     "For the soybean assemblies, the FASTA files are large (1 GB+) so server path is recommended.")
doc.add_paragraph()

# 5.4 Merge
heading(doc, "5.4  Merge VCFs", 2)
body(doc,
     "Combines two or more VCF files across samples at shared variant sites. "
     "Wraps bcftools merge with optional pre-normalization.")
doc.add_paragraph()

merge_opts = [
    ("Normalize multiallelic sites", "Runs bcftools norm -m +any on each input before merging", "Recommended when combining data from different pipelines"),
    ("Sample rename mapping",        "Upload a TSV of old_name → new_name pairs",               "Applied via bcftools reheader before merge"),
    ("Output filename",              "User-specified filename for the merged output",             "Saved as .vcf.gz (bgzipped)"),
]
simple_table(doc, ["Option", "Description", "Notes"], merge_opts, header_color="1F497D")
doc.add_paragraph()

body(doc,
     "Inputs must be indexed (.csi or .tbi). The backend automatically runs bcftools index "
     "on any unindexed input before merging. The result card shows total sites, total "
     "samples, and any merge warnings from stderr.")
doc.add_paragraph()

# 5.5 VCF Generator
heading(doc, "5.5  VCF Generator", 2)
body(doc,
     "Converts genotype dosage matrices (the format produced by Agriplex imputation and "
     "Gencove low-pass imputation) into standard VCF files. This is a Python port of the "
     "core logic in AGP_PYT_2025_MasterSet_GenoData_Prep.R.")
doc.add_paragraph()

heading(doc, "Input files", 3)
inputs_gen = [
    ("Genotype matrix", ".genotypes / .csv / .tsv / .txt",
     "Rows = samples OR SNPs (auto-detected). Values are dosage codes: 0, 1, 2, NA."),
    ("SNP info table",  ".txt / .csv / .tsv / .bed",
     "Columns: CHROM, POS, REF, ALT (case-insensitive, flexible aliases). ID auto-generated as CHROM-POS if absent."),
]
simple_table(doc, ["Input", "Formats", "Notes"], inputs_gen, header_color="1F6B3E")
doc.add_paragraph()

heading(doc, "Dosage translation", 3)
trans_rows = [
    ("0",          "0/0",  "Homozygous reference"),
    ("1",          "0/1",  "Heterozygous"),
    ("2",          "1/1",  "Homozygous alternate"),
    ("NA / empty", "./.",  "Missing genotype"),
]
simple_table(doc, ["Input dosage", "VCF GT", "Meaning"], trans_rows, header_color="555555")
doc.add_paragraph()

heading(doc, "Assembly selection", 3)
body(doc, "Supported assemblies for the ##reference header line:")
for asm in ["Wm82.a6.v1", "Wm82.a4.v1", "Wm82.a1.v1", "GRCh38", "GRCh37", "Custom (free text input)"]:
    bullet(doc, asm)
body(doc,
     "Variants are sorted by natural chromosome order (Gm01 < Gm02 < … < Gm20 for soybean). "
     "The result card reports sample count, SNP count, and missing genotype percentage.")
doc.add_paragraph()

# ── 6. SLURM / HPC Integration ───────────────────────────────────────────────

heading(doc, "6. SLURM / HPC Integration (MSI)", 1, "1F497D")
body(doc,
     "For large VCF operations (hundreds of samples, 50K+ variants) that exceed the "
     "capacity of the local web server, VariantTools can generate ready-to-run SLURM "
     "batch scripts for submission on the Minnesota Supercomputing Institute (MSI).")
doc.add_paragraph()

heading(doc, "6.1  Script Generation", 2)
body(doc,
     "After a Merge or Liftover job completes, an 'HPC / SLURM' section appears with a "
     "'Download SLURM Script' button. The generated script:")
for item in [
    "Uses the exact bcftools / CrossMap commands from the validated lab workflows",
    "Includes #SBATCH directives: --partition=amdsmall, --cpus-per-task, --mem, --time",
    "Loads the required environment modules: module load bcftools, module load crossmap",
    "Is pre-configured with the input/output file paths from the current job",
]:
    bullet(doc, item)
body(doc, "The user submits the script on MSI:  sbatch <script>.sh")
doc.add_paragraph()

heading(doc, "6.2  Configuration", 2)
body(doc, "SLURM settings are controlled via environment variables in docker-compose.yml:")
slurm_vars = [
    ("SLURM_ENABLED",   "false",                        "Set to 'true' to enable SLURM script generation"),
    ("SLURM_HOST",      "login.msi.umn.edu",            "MSI login node hostname"),
    ("SLURM_USER",      "(empty)",                       "MSI username (e.g. vramasub)"),
    ("SLURM_WORK_DIR",  "/scratch.global/varianttools",  "Working directory on MSI scratch"),
    ("SLURM_PARTITION", "amdsmall",                      "SLURM partition name"),
    ("SLURM_ACCOUNT",   "(empty)",                       "SLURM account/group if required"),
]
simple_table(doc, ["Variable", "Default", "Description"], slurm_vars, header_color="4B4B4B")
doc.add_paragraph()

# ── 7. Deployment ─────────────────────────────────────────────────────────────

heading(doc, "7. Deployment", 1, "1F497D")

heading(doc, "7.1  Local / Lab Server (Docker Compose)", 2)
body(doc, "Requirements: Docker Desktop, 8+ GB RAM, bcftools available inside container (pre-installed).")
code_block(doc,
"""# Start all services
cd VariantTools
docker-compose up --build

# Frontend:  http://localhost:3000
# API docs:  http://localhost:8000/docs
# Stop:      docker-compose down""")
doc.add_paragraph()

body(doc, "Four Docker services run together:")
svcs = [
    ("varianttools-redis-1",        "Redis 7 Alpine",       "Message broker and Celery result backend"),
    ("varianttools-backend-1",      "Python 3.11 + bcftools","FastAPI API server on port 8000"),
    ("varianttools-celery_worker-1","Python 3.11 + bcftools","Celery worker (concurrency=2)"),
    ("varianttools-frontend-1",     "Node 20 Alpine",        "Next.js production server on port 3000"),
]
simple_table(doc, ["Container", "Image", "Role"], svcs)
doc.add_paragraph()

heading(doc, "7.2  MSI / HPC (Singularity)", 2)
body(doc,
     "The Singularity definition file at containers/varianttools.def packages the FastAPI "
     "backend and all bioinformatics dependencies into a single portable image that can "
     "run on MSI without root access.")
code_block(doc,
"""# Build the image (run once on MSI or local Linux)
singularity build varianttools.sif containers/varianttools.def

# Start the API server
singularity run --bind ./data:/app/data varianttools.sif

# Start a Celery worker (separate terminal)
singularity exec --bind ./data:/app/data varianttools.sif \\
    celery -A backend.tasks.celery_app worker --loglevel=info

# Redis must be running separately (or use a lab server Redis)
# module load redis; redis-server --daemonize yes""")
doc.add_paragraph()

heading(doc, "7.3  Data Persistence", 2)
body(doc,
     "All runtime data is stored in the data/ directory, which is mounted as a volume "
     "in Docker and as a bind mount in Singularity. This directory is NOT tracked by git.")
data_dirs = [
    ("data/uploads/",         "Uploaded input files (UUID-prefixed, retained until cleanup)"),
    ("data/outputs/",         "Generated output files (VCFs, stats, merged files)"),
    ("data/chains/",          "Chain files for liftover (place .chain files here manually)"),
    ("data/varianttools.db",  "SQLite database — job history and chain file registry"),
]
simple_table(doc, ["Path", "Contents"], data_dirs)
doc.add_paragraph()

# ── 8. Original Workflow Reference ───────────────────────────────────────────

heading(doc, "8. Original Workflow Scripts", 1, "1F497D")
body(doc,
     "The app was built by systematically porting the following scripts. "
     "They remain in the repository as reference and documentation.")

scripts = [
    ("AGP_PYT_2025_MasterSet_GenoData_Prep.R",
     "R",
     "Master genotyping data prep for PYT 2025. Source of VCF Generator and fixref logic. "
     "Covers Agriplex .genotypes loading, dosage translation, A4→A6 coordinate merging, "
     "VCF writing, bcftools annotate/stats/fixref, and multi-year PYT merging."),
    ("Liftover_a4_to_a6.txt",
     "Shell + R",
     "SoySNP50K liftover from Wm82.a4 to Wm82.a6. Documents MUMmer alignment, "
     "CrossMap and UCSC liftOver workflows, minimap2/paftools comparison, and "
     "R post-processing to build the A4→A6 cross-map table (36,164 SNPs)."),
    ("Liftover_A1_toA6.txt",
     "Shell + R",
     "AGP SoySNP1K liftover from Wm82.a1 to Wm82.a6. Same pipeline as A4→A6, "
     "yielding 1,242 uniquely mapped positions."),
    ("Gencove_Subsets_6K_50K_bcfTools.txt",
     "Shell + R",
     "Gencove low-pass data management: S3 download, BARCSoySNP6K/50K/1K subsetting, "
     "sample renaming, multiallelic normalisation, multi-year merge (2020+2022), "
     "chromosome prefix reformatting, and GATK validation."),
]
simple_table(doc, ["File", "Language", "Contents"], scripts, header_color="1F497D")
doc.add_paragraph()

# ── 9. API Reference ──────────────────────────────────────────────────────────

heading(doc, "9. Full API Reference", 1, "1F497D")
body(doc, "All endpoints are documented interactively at http://localhost:8000/docs (Swagger UI).")
doc.add_paragraph()

all_routes = [
    # stats
    ("POST", "/api/stats/upload-and-analyze",     "Upload VCF; returns assembly_info + job_id"),
    ("GET",  "/api/stats/result/{job_id}",         "Poll stats job result"),
    ("GET",  "/api/stats/download/{job_id}",       "Download raw stats text"),
    # liftover
    ("GET",  "/api/liftover/chain-files",          "List registered chain files"),
    ("POST", "/api/liftover/chain-files/upload",   "Upload and register a chain file"),
    ("POST", "/api/liftover/upload",               "Upload VCF or BED input"),
    ("POST", "/api/liftover/submit",               "Submit liftover job"),
    ("GET",  "/api/liftover/status/{job_id}",      "Poll liftover job result"),
    ("GET",  "/api/liftover/download/{job_id}",    "Download mapped output"),
    ("GET",  "/api/liftover/download-unmapped/{job_id}", "Download unmapped entries"),
    # fixref
    ("POST", "/api/fixref/upload",                 "Upload VCF for fixref"),
    ("POST", "/api/fixref/upload-ref",             "Upload reference FASTA"),
    ("POST", "/api/fixref/check",                  "Run check mode (synchronous, fast)"),
    ("POST", "/api/fixref/fix",                    "Submit fix job (async)"),
    ("GET",  "/api/fixref/status/{job_id}",        "Poll fix job result"),
    ("GET",  "/api/fixref/download/{job_id}",      "Download fixed VCF"),
    # merge
    ("POST", "/api/merge/upload",                  "Upload one or more VCF files"),
    ("POST", "/api/merge/submit",                  "Submit merge job with options"),
    ("GET",  "/api/merge/status/{job_id}",         "Poll merge job result"),
    ("GET",  "/api/merge/download/{job_id}",       "Download merged VCF"),
    # generator
    ("POST", "/api/generator/upload-genotypes",    "Upload dosage matrix; returns preview"),
    ("POST", "/api/generator/upload-snp-info",     "Upload SNP metadata table; returns preview"),
    ("POST", "/api/generator/generate",            "Submit VCF generation job"),
    ("GET",  "/api/generator/status/{job_id}",     "Poll generator job result"),
    ("GET",  "/api/generator/download/{job_id}",   "Download generated VCF"),
    # slurm
    ("GET",  "/api/slurm/status",                  "Return SLURM configuration status"),
    ("GET",  "/api/slurm/script/merge",            "Download SLURM script for a merge job"),
    ("GET",  "/api/slurm/script/liftover",         "Download SLURM script for a liftover job"),
    ("GET",  "/api/slurm/script/fixref",           "Download SLURM script for a fixref job"),
    # health
    ("GET",  "/health",                            "API liveness probe"),
]
simple_table(doc, ["Method", "Endpoint", "Description"], all_routes, header_color="2E75B6")
doc.add_paragraph()

# ── 10. Future Work ───────────────────────────────────────────────────────────

heading(doc, "10. Planned Future Modules", 1, "1F497D")
future = [
    ("Marker Panel Subsetting", "bcftools view -T for BARCSoySNP6K / SoySNP50K / custom panel BED files"),
    ("Sample ID Manager",       "bcftools reheader -s with UI table for mapping old→new names; duplicate detection"),
    ("VCF Annotator",           "SnpEff / SnpSift integration for functional annotation"),
    ("SLURM Auto-Submit",       "SSH-based job submission to MSI (requires key-based auth configuration)"),
    ("S3 Integration",          "Direct s3cmd get/put for Gencove merged VCFs from lorenzgenotyping bucket"),
    ("Genotype Explorer",       "Interactive PCA and kinship matrix from VCF data"),
]
simple_table(doc, ["Module", "Description"], future)
doc.add_paragraph()

# ── Footer ────────────────────────────────────────────────────────────────────

add_horizontal_rule(doc)
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(
    f"VariantTools v0.3.0  ·  Lorenz Lab, University of Minnesota  ·  "
    f"Generated {datetime.date.today().strftime('%B %d, %Y')}"
)
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── Save ──────────────────────────────────────────────────────────────────────

out_path = "VariantTools_Architecture_Guide.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
